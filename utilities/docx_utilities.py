from docx.shared import Pt
import requests
from docx.shared import Inches
from io import BytesIO
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
# microsoft document formation logic 

# Iterate through the JSON data and add content to the document
def Form_doc(data, doc_name, level=0):
    

    def process_item(item, parent,level):
        text = item.get('text', '')
        block_type = item.get('type', '')
        children = item.get('children', [])
        table_data = item.get('table_data', [])
        database_data = item.get('database_data', [])
        image_url = item.get('image_url', '')

        if block_type == 'heading_1':
            Add_paragraph(parent, text, style='Heading 1', level=level)
        elif block_type == 'heading_2':
            Add_paragraph(parent, text, style='Heading 2', level=level)
        elif block_type == 'heading_3':
            Add_paragraph(parent, text, style='Heading 3', level=level)        
        elif block_type == 'paragraph':
            Add_paragraph(parent, text, level=level)
        elif block_type == 'bulleted_list_item':
            Add_paragraph(parent, f" {text}", style='List Bullet', level=level)
        elif block_type == 'to_do':
            Add_paragraph(parent, f"[ ] {text}", style='List Bullet', level=level)
        elif block_type == 'table':
            Add_table(parent, table_data)
        elif block_type == 'child_table':
            Add_table(parent, table_data)        
        elif block_type == 'child_database':
            Add_database(parent, database_data)
        elif block_type == 'image':
            Add_image_from_url(parent, image_url)

        # Recursively process children
        if children:
            for child in children:
                process_item(child, parent, level + 3)
    
    doc = Document()

    for item in data:
        process_item(item, doc, level)
    
    doc.save(f'local/{doc_name}.docx')
    # print(f'local/{doc_name}.docx')

def Add_paragraph(doc, text, style=None, bold=False, underline=False, alignment=None, level=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if style:
        p.style = style
    if bold:
        run.bold = True
    if underline:
        run.underline = True
    if alignment:
        p.alignment = alignment

    # Indentation for nested content
    if level > 0:
        p.paragraph_format.left_indent = Pt(14 * level)

def Add_table(doc, table_data):
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            row.cells[col_idx].text = cell_text

def Add_database(doc, table_data):
    # Create a table with the correct number of rows and columns
    table = doc.add_table(rows=len(table_data) + 1, cols=len(table_data[0]))

    # Add the headers
    headers = list(table_data[0].keys())
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header

    # Add the data rows
    for row_idx, row_data in enumerate(table_data, start=1):
        for col_idx, (key, value) in enumerate(row_data.items()):
            table.cell(row_idx, col_idx).text = str(value)

def Add_image_from_url(doc, url, width=None, height=None):
    try:
        # Send request to fetch the image
        response = requests.get(url)
        
        # Check if the response status is 200 (OK)
        if response.status_code != 200:
            print(f"Error fetching image from URL. Status code: {response.status_code}")
        
        # Check content type to ensure it's an image
        content_type = response.headers['Content-Type']
        if not content_type.startswith("image"):
            print(f"URL does not point to a valid image. Content-Type: {content_type}")
        
        # Convert image content into BytesIO stream
        image_stream = BytesIO(response.content)
        
        # Add the image to the document
        if width and height:
            doc.add_picture(image_stream, width=Inches(width), height=Inches(height))
        else:
            doc.add_picture(image_stream)
    
    except UnrecognizedImageError:
        print(f"The image at {url} is not in a recognized format.")
    
    except ValueError as e:
        # Handle any other errors (network issues, invalid URL, etc.)
        print(f"An error occurred while adding image from URL: {str(e)} : {url}")